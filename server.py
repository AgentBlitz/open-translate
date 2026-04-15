import io
import os
import re
import threading
from typing import Any, Dict, List, Optional, Tuple, Union

import numpy as np
import torch
import deepspeed
from fastapi import FastAPI, File, Form, HTTPException, Query, Request, UploadFile
from fastapi.responses import StreamingResponse
from fastapi.staticfiles import StaticFiles
from PIL import Image
from pydantic import BaseModel, Field
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM

import pycountry
from langdetect import detect as ld_detect
from langdetect.lang_detect_exception import LangDetectException

Image.MAX_IMAGE_PIXELS = 50_000_000


SIZE_TO_REPO = {
    "600M": "facebook/nllb-200-distilled-600M",
    "600M-distilled": "facebook/nllb-200-distilled-600M",
    "1.3B": "facebook/nllb-200-1.3B",
    "1.3B-distilled": "facebook/nllb-200-distilled-1.3B",
    "3.3B": "facebook/nllb-200-3.3B",
}


def pick_model_id() -> str:
    override = os.getenv("NLLB_MODEL_ID", "").strip()
    if override:
        return override
    size = os.getenv("NLLB_MODEL_SIZE", "600M").strip()
    if size not in SIZE_TO_REPO:
        raise ValueError(
            f"Unsupported NLLB_MODEL_SIZE={size}. Use one of {list(SIZE_TO_REPO.keys())} or set NLLB_MODEL_ID."
        )
    return SIZE_TO_REPO[size]


def pick_dtype() -> torch.dtype:
    dt = os.getenv("DTYPE", "fp16").strip().lower()
    if dt in ("fp16", "float16"):
        return torch.float16
    if dt in ("bf16", "bfloat16"):
        return torch.bfloat16
    if dt in ("fp32", "float32"):
        return torch.float32
    raise ValueError("DTYPE must be one of: fp16, bf16, fp32")


def pick_tp_size() -> int:
    tp = os.getenv("TP_SIZE", "auto").strip().lower()
    if tp == "auto":
        return max(1, torch.cuda.device_count() if torch.cuda.is_available() else 1)
    return max(1, int(tp))


class TranslateRequest(BaseModel):
    q: Union[str, List[str]] = Field(
        ..., description="Text or list of texts to translate."
    )
    target: str = Field(
        ..., description="Target language (ISO 639-1 or BCP-47 like 'es', 'zh-TW')."
    )
    source: Optional[str] = Field(
        None,
        description="Optional source language (ISO 639-1 or BCP-47). If omitted, best-effort detect.",
    )
    format: Optional[str] = Field(
        "text", description="Compatibility field; 'text' or 'html'."
    )
    model: Optional[str] = Field(None, description="Compatibility field; ignored.")
    max_new_tokens: int = Field(128, ge=1, le=2048)
    num_beams: int = Field(1, ge=1, le=16)
    temperature: float = Field(1.0, gt=0.0, le=5.0)
    top_p: float = Field(1.0, gt=0.0, le=1.0)
    do_sample: bool = Field(False)
    truncate_input: bool = Field(True)


app = FastAPI(title="NLLB-200 Inference API")


@app.middleware("http")
async def _no_store_headers(request: Request, call_next):
    response = await call_next(request)
    response.headers["Cache-Control"] = "no-store, no-cache, must-revalidate"
    response.headers["Pragma"] = "no-cache"
    response.headers["Referrer-Policy"] = "no-referrer"
    return response


_UI_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "static", "ui")
if os.path.isdir(_UI_DIR):
    app.mount("/ui", StaticFiles(directory=_UI_DIR, html=True), name="ui")

MODEL_ID = pick_model_id()
DTYPE = pick_dtype()
MAX_BATCH_SIZE = int(os.getenv("MAX_BATCH_SIZE", "32"))
MAX_INPUT_LENGTH = int(os.getenv("MAX_INPUT_LENGTH", "1024"))

MAX_DOC_BYTES = int(os.getenv("MAX_DOC_BYTES", str(25 * 1024 * 1024)))
MAX_PDF_PAGES = int(os.getenv("MAX_PDF_PAGES", "50"))
OCR_DPI = int(os.getenv("OCR_DPI", "200"))
OCR_LANG = os.getenv("OCR_LANG", "en")

tokenizer = None
model = None
ds_engine = None
_MODEL_LOCK = threading.Lock()

_ocr_engine = None
_OCR_LOCK = threading.Lock()

ISO_TO_NLLB: Dict[str, str] = {}
ISO_NAME: Dict[str, str] = {}

# Display names for NLLB ISO-639-3 codes that pycountry doesn't cover well
# (Arabic dialects, regional languages, NLLB-specific codes). Keyed by iso3.
NLLB_NAME_OVERRIDES: Dict[str, str] = {
    "ace": "Acehnese",
    "acm": "Mesopotamian Arabic",
    "acq": "Ta'izzi-Adeni Arabic",
    "aeb": "Tunisian Arabic",
    "ajp": "South Levantine Arabic",
    "apc": "North Levantine Arabic",
    "arb": "Modern Standard Arabic",
    "ars": "Najdi Arabic",
    "ary": "Moroccan Arabic",
    "arz": "Egyptian Arabic",
    "awa": "Awadhi",
    "ayr": "Central Aymara",
    "azb": "South Azerbaijani",
    "azj": "North Azerbaijani",
    "bam": "Bambara",
    "ban": "Balinese",
    "bem": "Bemba",
    "bho": "Bhojpuri",
    "bjn": "Banjar",
    "bug": "Buginese",
    "ceb": "Cebuano",
    "cjk": "Chokwe",
    "crh": "Crimean Tatar",
    "dik": "Southwestern Dinka",
    "dyu": "Dyula",
    "dzo": "Dzongkha",
    "ewe": "Ewe",
    "fij": "Fijian",
    "fon": "Fon",
    "fur": "Friulian",
    "fuv": "Nigerian Fulfulde",
    "gaz": "West Central Oromo",
    "gla": "Scottish Gaelic",
    "grn": "Guarani",
    "hat": "Haitian Creole",
    "hne": "Chhattisgarhi",
    "ibo": "Igbo",
    "ilo": "Ilocano",
    "kab": "Kabyle",
    "kac": "Jingpho",
    "kam": "Kamba",
    "kas": "Kashmiri",
    "kbp": "Kabiyè",
    "kea": "Kabuverdianu",
    "khk": "Halh Mongolian",
    "kik": "Kikuyu",
    "kin": "Kinyarwanda",
    "kmb": "Kimbundu",
    "kmr": "Northern Kurdish",
    "knc": "Central Kanuri",
    "kon": "Kikongo",
    "lij": "Ligurian",
    "lim": "Limburgish",
    "lin": "Lingala",
    "lmo": "Lombard",
    "ltg": "Latgalian",
    "lua": "Luba-Kasai",
    "lug": "Ganda",
    "luo": "Luo",
    "lus": "Mizo",
    "mag": "Magahi",
    "mai": "Maithili",
    "min": "Minangkabau",
    "mni": "Meitei",
    "mos": "Mossi",
    "nqo": "N'Ko",
    "nso": "Northern Sotho",
    "nus": "Nuer",
    "nya": "Nyanja",
    "oci": "Occitan",
    "ory": "Odia",
    "pag": "Pangasinan",
    "pap": "Papiamento",
    "pbt": "Southern Pashto",
    "pes": "Western Persian",
    "plt": "Plateau Malagasy",
    "prs": "Dari",
    "quy": "Ayacucho Quechua",
    "run": "Rundi",
    "sag": "Sango",
    "san": "Sanskrit",
    "sat": "Santali",
    "scn": "Sicilian",
    "shn": "Shan",
    "sna": "Shona",
    "snd": "Sindhi",
    "som": "Somali",
    "sot": "Southern Sotho",
    "srd": "Sardinian",
    "ssw": "Swati",
    "sun": "Sundanese",
    "swh": "Swahili",
    "szl": "Silesian",
    "tam": "Tamil",
    "taq": "Tamasheq",
    "tgk": "Tajik",
    "tir": "Tigrinya",
    "tpi": "Tok Pisin",
    "tsn": "Tswana",
    "tso": "Tsonga",
    "tuk": "Turkmen",
    "tum": "Tumbuka",
    "twi": "Twi",
    "tzm": "Central Atlas Tamazight",
    "uig": "Uyghur",
    "umb": "Umbundu",
    "vec": "Venetian",
    "war": "Waray",
    "wol": "Wolof",
    "xho": "Xhosa",
    "ydd": "Eastern Yiddish",
    "yor": "Yoruba",
    "yue": "Cantonese",
    "zsm": "Standard Malay",
    "zul": "Zulu",
    "bak": "Bashkir",
    "kaz": "Kazakh",
    "kir": "Kyrgyz",
    "uzn": "Northern Uzbek",
    "lao": "Lao",
    "mya": "Burmese",
    "khm": "Khmer",
    "tha": "Thai",
    "vie": "Vietnamese",
    "jpn": "Japanese",
    "kor": "Korean",
    "heb": "Hebrew",
    "ell": "Greek",
    "hrv": "Croatian",
    "srp": "Serbian",
    "bos": "Bosnian",
    "mkd": "Macedonian",
    "slv": "Slovenian",
    "ckb": "Central Kurdish",
}

BCP47_SPECIALS = {
    "zh-cn": "zho_Hans",
    "zh-sg": "zho_Hans",
    "zh-hans": "zho_Hans",
    "zh-tw": "zho_Hant",
    "zh-hk": "zho_Hant",
    "zh-mo": "zho_Hant",
    "zh-hant": "zho_Hant",
}


def _pycountry_name(iso: str) -> Optional[str]:
    try:
        if len(iso) == 2:
            lang = pycountry.languages.get(alpha_2=iso)
        else:
            lang = pycountry.languages.get(alpha_3=iso)
        return getattr(lang, "name", None) if lang else None
    except Exception:
        return None


def _display_name(iso: str) -> Optional[str]:
    # Prefer our curated NLLB name overrides, then pycountry.
    if iso in NLLB_NAME_OVERRIDES:
        return NLLB_NAME_OVERRIDES[iso]
    return _pycountry_name(iso)


def _iso1_from_iso3(iso3: str) -> Optional[str]:
    try:
        lang = pycountry.languages.get(alpha_3=iso3)
        return getattr(lang, "alpha_2", None) if lang else None
    except Exception:
        return None


def _normalize_bcp47(code: str) -> str:
    return code.strip().replace("_", "-").lower()


def _preferred_variant(nllb_codes: List[str]) -> str:
    # Prefer Latn variants first
    for s in ("Latn", "Cyrl", "Arab", "Deva", "Hans", "Hant"):
        for c in nllb_codes:
            if c.endswith("_" + s):
                return c
    return nllb_codes[0]


def _build_language_maps():
    global ISO_TO_NLLB, ISO_NAME

    lang_code_to_id = _get_lang_code_to_id()
    codes = sorted(lang_code_to_id.keys())
    by_iso3: Dict[str, List[str]] = {}
    for c in codes:
        if "_" not in c:
            continue
        iso3, _script = c.split("_", 1)
        by_iso3.setdefault(iso3.lower(), []).append(c)

    ISO_TO_NLLB = {}
    ISO_NAME = {}

    for iso3, variants in by_iso3.items():
        chosen = _preferred_variant(sorted(variants))
        iso1 = _iso1_from_iso3(iso3)

        ISO_TO_NLLB[iso3] = chosen
        ISO_NAME.setdefault(iso3, _display_name(iso3) or iso3)

        if iso1:
            iso1 = iso1.lower()
            ISO_TO_NLLB[iso1] = chosen
            ISO_NAME.setdefault(iso1, _display_name(iso1) or _display_name(iso3) or iso1)

    # BCP-47 region/script special-cases
    for k, v in BCP47_SPECIALS.items():
        if v in lang_code_to_id:
            ISO_TO_NLLB[k] = v


def _resolve_to_nllb(code: str) -> str:
    if not code:
        raise HTTPException(status_code=400, detail="Missing language code.")
    norm = _normalize_bcp47(code)

    # Exact match
    if norm in ISO_TO_NLLB:
        return ISO_TO_NLLB[norm]

    # Language-only fallback
    lang_part = norm.split("-", 1)[0]
    if lang_part in ISO_TO_NLLB:
        return ISO_TO_NLLB[lang_part]

    raise HTTPException(status_code=400, detail=f"Unsupported language code: '{code}'")


def _lang_to_id(nllb_tgt: str) -> int:
    lang_code_to_id = _get_lang_code_to_id()
    if nllb_tgt not in lang_code_to_id:
        raise HTTPException(
            status_code=400, detail=f"Unknown NLLB target code '{nllb_tgt}'"
        )
    return int(lang_code_to_id[nllb_tgt])


_NLLB_LANG_RE = re.compile(r"^[a-z]{3}_[A-Za-z]{4}$")


def _get_lang_code_to_id() -> Dict[str, int]:
    if hasattr(tokenizer, "lang_code_to_id"):
        return tokenizer.lang_code_to_id
    vocab = tokenizer.get_vocab()
    return {tok: idx for tok, idx in vocab.items() if _NLLB_LANG_RE.match(tok)}


def _detect_iso639_1(text: str) -> Optional[str]:
    try:
        return ld_detect(text).lower()
    except LangDetectException:
        return None
    except Exception:
        return None


def _ensure_batch(q: Union[str, List[str]]) -> List[str]:
    return [q] if isinstance(q, str) else q


def _detect_one(text: str) -> Dict[str, Any]:
    iso = _detect_iso639_1(text) or "en"
    return {"language": iso}


@app.on_event("startup")
def _startup():
    global tokenizer, model, ds_engine, _ocr_engine

    if torch.cuda.is_available():
        torch.cuda.empty_cache()

    if torch.cuda.is_available():
        try:
            deepspeed.init_distributed()
        except Exception:
            pass

    tokenizer = AutoTokenizer.from_pretrained(MODEL_ID, use_fast=True)
    model = AutoModelForSeq2SeqLM.from_pretrained(
        MODEL_ID,
        dtype=DTYPE,
        low_cpu_mem_usage=True,
    )
    model.eval()

    if torch.cuda.is_available():
        tp_size = pick_tp_size()
        try:
            ds_engine = deepspeed.init_inference(
                model=model,
                dtype=DTYPE,
                replace_with_kernel_inject=True,
                tensor_parallel={"tp_size": tp_size},
            )
        except Exception as e:
            print(
                f"[startup] DeepSpeed kernel injection unavailable ({type(e).__name__}): "
                f"{e}. Falling back to plain HF model on CUDA.",
                flush=True,
            )
            ds_engine = None
            model.to("cuda")
    else:
        ds_engine = None

    _build_language_maps()

    try:
        from paddleocr import PaddleOCR

        device = "gpu" if torch.cuda.is_available() else "cpu"
        _ocr_engine = PaddleOCR(
            use_doc_orientation_classify=False,
            use_doc_unwarping=False,
            use_textline_orientation=True,
            lang=OCR_LANG,
            device=device,
        )
        print(f"[startup] PaddleOCR ready on {device} (lang={OCR_LANG})", flush=True)
    except Exception as e:
        print(
            f"[startup] PaddleOCR unavailable ({type(e).__name__}): {e}. "
            f"Document translation will reject image / scanned-PDF inputs.",
            flush=True,
        )
        _ocr_engine = None


@torch.inference_mode()
def _translate_batch(
    texts: List[str], src_nllb: str, tgt_nllb: str, gen: Dict[str, Any]
) -> List[str]:
    if len(texts) > MAX_BATCH_SIZE:
        raise HTTPException(
            status_code=400,
            detail=f"Batch too large: {len(texts)} > MAX_BATCH_SIZE={MAX_BATCH_SIZE}",
        )

    forced_bos_token_id = _lang_to_id(tgt_nllb)

    with _MODEL_LOCK:
        tokenizer.src_lang = src_nllb
        enc = tokenizer(
            texts,
            return_tensors="pt",
            padding=True,
            truncation=gen.get("truncate_input", True),
            max_length=MAX_INPUT_LENGTH,
        )
        if torch.cuda.is_available():
            enc = {k: v.to("cuda") for k, v in enc.items()}

        gen_model = ds_engine.module if ds_engine is not None else model
        generated = gen_model.generate(
            **enc,
            forced_bos_token_id=forced_bos_token_id,
            max_new_tokens=gen.get("max_new_tokens", 128),
            num_beams=gen.get("num_beams", 1),
            do_sample=gen.get("do_sample", False),
            temperature=gen.get("temperature", 1.0),
            top_p=gen.get("top_p", 1.0),
        )
        return tokenizer.batch_decode(generated, skip_special_tokens=True)


@app.get("/health")
def health() -> Dict[str, Any]:
    return {
        "status": "ok",
        "model_id": MODEL_ID,
        "cuda": torch.cuda.is_available(),
        "gpus": torch.cuda.device_count() if torch.cuda.is_available() else 0,
        "dtype": str(DTYPE),
        "max_batch_size": MAX_BATCH_SIZE,
        "max_input_length": MAX_INPUT_LENGTH,
    }


@app.post("/language/translate/v2")
def translate(req: TranslateRequest) -> Dict[str, Any]:
    if tokenizer is None or model is None:
        raise HTTPException(status_code=503, detail="Model not loaded yet.")

    texts = _ensure_batch(req.q)
    if not texts:
        raise HTTPException(status_code=400, detail="Missing q")

    tgt_nllb = _resolve_to_nllb(req.target)

    detected_iso: Optional[str] = None
    if req.source:
        src_nllb = _resolve_to_nllb(req.source)
    else:
        detected_iso = _detect_iso639_1(texts[0]) or "en"
        try:
            src_nllb = _resolve_to_nllb(detected_iso)
        except HTTPException:
            detected_iso = "en"
            src_nllb = _resolve_to_nllb("en")

    out = _translate_batch(
        texts,
        src_nllb=src_nllb,
        tgt_nllb=tgt_nllb,
        gen=req.model_dump(),
    )

    translations = []
    for s in out:
        item = {"translatedText": s}
        if req.source is None:
            item["detectedSourceLanguage"] = detected_iso
        translations.append(item)

    return {"data": {"translations": translations}}


class DetectRequest(BaseModel):
    q: Union[str, List[str]] = Field(
        ..., description="Text or list of texts to detect."
    )


@app.post("/language/translate/v2/detect")
def detect(req: DetectRequest) -> Dict[str, Any]:
    texts = _ensure_batch(req.q)
    if not texts:
        raise HTTPException(status_code=400, detail="Missing q")

    detections = [[_detect_one(t)] for t in texts]
    return {"data": {"detections": detections}}


@app.get("/language/translate/v2/languages")
def languages(
    target: Optional[str] = Query(
        None,
        description="Optional target language for localized names (compatibility).",
    ),
) -> Dict[str, Any]:
    # Dedupe by NLLB target code so each physical language appears once.
    # Prefer the shortest ISO key (ISO-639-1 when available, else ISO-639-3).
    best_key_for_nllb: Dict[str, str] = {}
    for iso, nllb in ISO_TO_NLLB.items():
        if "-" in iso:
            continue
        current = best_key_for_nllb.get(nllb)
        if current is None or len(iso) < len(current):
            best_key_for_nllb[nllb] = iso

    langs = []
    for nllb, iso in best_key_for_nllb.items():
        entry = {"language": iso}
        if target:
            entry["name"] = ISO_NAME.get(iso, iso)
        langs.append(entry)

    langs.sort(key=lambda x: x.get("name", x["language"]).lower())
    return {"data": {"languages": langs}}


# ---------------------------------------------------------------------------
# Document translation: DOCX / PDF / images via PaddleOCR + NLLB
# ---------------------------------------------------------------------------

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
_IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".webp", ".bmp", ".tif", ".tiff"}


def _classify_upload(filename: str, content_type: str) -> str:
    name = (filename or "").lower()
    ct = (content_type or "").lower()
    if name.endswith(".docx") or ct == DOCX_MIME:
        return "docx"
    if name.endswith(".pdf") or ct == "application/pdf":
        return "pdf"
    if ct.startswith("image/") or any(name.endswith(ext) for ext in _IMAGE_EXTS):
        return "image"
    return "unknown"


def _extract_docx_runs(data: bytes):
    from docx import Document

    doc = Document(io.BytesIO(data))
    runs = []
    for para in doc.paragraphs:
        for run in para.runs:
            if run.text and run.text.strip():
                runs.append(run)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.text and run.text.strip():
                            runs.append(run)
    return doc, runs


def _extract_pdf_text_layer(data: bytes) -> List[str]:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages: List[str] = []
    total_chars = 0
    for page in reader.pages[:MAX_PDF_PAGES]:
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        pages.append(text)
        total_chars += len(text.strip())
    if total_chars < 20:
        return []
    return pages


def _rasterize_pdf(data: bytes, dpi: int) -> List[Image.Image]:
    import pypdfium2 as pdfium

    pdf = pdfium.PdfDocument(io.BytesIO(data))
    try:
        n = min(len(pdf), MAX_PDF_PAGES)
        scale = dpi / 72.0
        images: List[Image.Image] = []
        for i in range(n):
            page = pdf[i]
            try:
                pil = page.render(scale=scale).to_pil().convert("RGB")
                images.append(pil)
            finally:
                page.close()
        return images
    finally:
        pdf.close()


def _ocr_image(img: Image.Image) -> List[str]:
    if _ocr_engine is None:
        raise HTTPException(
            status_code=503,
            detail="OCR engine not loaded; image and scanned-PDF inputs are unavailable.",
        )
    arr = np.array(img.convert("RGB"))
    with _OCR_LOCK:
        result = _ocr_engine.predict(arr)
    lines: List[str] = []
    if not result:
        return lines
    for res in result:
        texts = None
        if hasattr(res, "get"):
            texts = res.get("rec_texts")
        if texts is None and hasattr(res, "__getitem__"):
            try:
                texts = res["rec_texts"]
            except Exception:
                texts = None
        if texts:
            for t in texts:
                if t and t.strip():
                    lines.append(t.strip())
    return lines


def _ocr_pdf_pages(pages: List[Image.Image]) -> List[str]:
    out: List[str] = []
    for i, page in enumerate(pages):
        if i > 0:
            out.append("")
        out.extend(_ocr_image(page))
    return out


def _translate_paragraphs(
    texts: List[str], src_nllb: str, tgt_nllb: str
) -> List[str]:
    if not texts:
        return []
    gen = {"max_new_tokens": 256, "num_beams": 1, "truncate_input": True}
    out: List[str] = []
    # Preserve empty-string separators without sending them through the model.
    idx_to_text: List[Tuple[int, str]] = [
        (i, t) for i, t in enumerate(texts) if t
    ]
    translated: List[Optional[str]] = [None] * len(texts)
    for blank_i, _ in [(i, t) for i, t in enumerate(texts) if not t]:
        translated[blank_i] = ""

    for start in range(0, len(idx_to_text), MAX_BATCH_SIZE):
        chunk = idx_to_text[start : start + MAX_BATCH_SIZE]
        chunk_texts = [t for _, t in chunk]
        chunk_out = _translate_batch(
            chunk_texts, src_nllb=src_nllb, tgt_nllb=tgt_nllb, gen=gen
        )
        for (orig_i, _), t_out in zip(chunk, chunk_out):
            translated[orig_i] = t_out

    for t in translated:
        out.append(t if t is not None else "")
    return out


def _build_translated_docx(paragraphs: List[str]) -> bytes:
    from docx import Document

    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _resolve_src_for_text(
    sample_text: str, source: Optional[str]
) -> str:
    if source:
        return _resolve_to_nllb(source)
    detected = _detect_iso639_1(sample_text) or "en"
    try:
        return _resolve_to_nllb(detected)
    except HTTPException:
        return _resolve_to_nllb("en")


@app.post("/language/translate/v2/document")
async def translate_document(
    file: UploadFile = File(...),
    target: str = Form(...),
    source: Optional[str] = Form(None),
) -> StreamingResponse:
    if tokenizer is None or model is None:
        raise HTTPException(status_code=503, detail="Model not loaded yet.")

    if file.size is not None and file.size > MAX_DOC_BYTES:
        raise HTTPException(status_code=413, detail="File too large.")

    data = await file.read()
    if not data:
        raise HTTPException(status_code=400, detail="Empty file.")
    if len(data) > MAX_DOC_BYTES:
        raise HTTPException(status_code=413, detail="File too large.")

    kind = _classify_upload(file.filename or "", file.content_type or "")
    tgt_nllb = _resolve_to_nllb(target)

    if kind == "docx":
        doc, runs = _extract_docx_runs(data)
        texts = [r.text for r in runs]
        if not texts:
            out_bytes = _build_translated_docx([])
        else:
            src_nllb = _resolve_src_for_text(texts[0], source)
            translated = _translate_paragraphs(texts, src_nllb, tgt_nllb)
            for run, new_text in zip(runs, translated):
                run.text = new_text
            buf = io.BytesIO()
            doc.save(buf)
            out_bytes = buf.getvalue()

    elif kind == "pdf":
        page_texts = _extract_pdf_text_layer(data)
        if page_texts:
            # Text-layer PDF: one paragraph per non-empty line, page-separated by blanks.
            paragraphs: List[str] = []
            for i, page_text in enumerate(page_texts):
                if i > 0:
                    paragraphs.append("")
                for line in page_text.splitlines():
                    line = line.strip()
                    if line:
                        paragraphs.append(line)
        else:
            images = _rasterize_pdf(data, OCR_DPI)
            paragraphs = _ocr_pdf_pages(images)
        if not paragraphs:
            raise HTTPException(status_code=422, detail="No text extracted from PDF.")
        sample = next((p for p in paragraphs if p), "")
        src_nllb = _resolve_src_for_text(sample, source)
        translated = _translate_paragraphs(paragraphs, src_nllb, tgt_nllb)
        out_bytes = _build_translated_docx(translated)

    elif kind == "image":
        try:
            img = Image.open(io.BytesIO(data))
            img.load()
        except Exception:
            raise HTTPException(status_code=400, detail="Unsupported or corrupt image.")
        paragraphs = _ocr_image(img)
        if not paragraphs:
            raise HTTPException(status_code=422, detail="No text detected in image.")
        sample = next((p for p in paragraphs if p), "")
        src_nllb = _resolve_src_for_text(sample, source)
        translated = _translate_paragraphs(paragraphs, src_nllb, tgt_nllb)
        out_bytes = _build_translated_docx(translated)

    else:
        raise HTTPException(
            status_code=415,
            detail="Unsupported file type. Use .docx, .pdf, or common image formats.",
        )

    return StreamingResponse(
        io.BytesIO(out_bytes),
        media_type=DOCX_MIME,
        headers={"Content-Disposition": 'attachment; filename="translated.docx"'},
    )
